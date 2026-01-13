from flask import Flask, render_template, request, send_file
import os, zipfile
from pypdf import PdfWriter, PdfReader
from pdf2docx import Converter
from PIL import Image
import pandas as pd
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = '/tmp' # Plus stable sur Render

@app.route('/')
def home():
    return render_template('index.html')

# --- PDF TOOLS ---
@app.route('/merge', methods=['POST'])
def merge():
    files = request.files.getlist("pdfs")
    merger = PdfWriter()
    for file in files: 
        merger.append(file)
    path = os.path.join(UPLOAD_FOLDER, "fusion.pdf")
    merger.write(path)
    merger.close()
    return send_file(path, as_attachment=True)

@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    f = request.files['file']
    in_p = os.path.join(UPLOAD_FOLDER, f.filename)
    out_p = in_p.replace('.pdf', '.docx')
    f.save(in_p)
    cv = Converter(in_p)
    cv.convert(out_p)
    cv.close()
    return send_file(out_p, as_attachment=True)

# --- CONVERTISSEURS SPÉCIAUX (LINUX COMPATIBLE) ---

# EXCEL VERS WORD
@app.route('/excel-to-word', methods=['POST'])
def excel_to_word():
    f = request.files['file']
    in_p = os.path.join(UPLOAD_FOLDER, f.filename)
    out_p = in_p.rsplit('.', 1)[0] + '.docx'
    f.save(in_p)
    df = pd.read_excel(in_p)
    doc = Document()
    doc.add_heading('Données Excel converties', 0)
    # Création d'un tableau dans Word
    t = doc.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    doc.save(out_p)
    return send_file(out_p, as_attachment=True)

# WORD VERS EXCEL
@app.route('/word-to-excel', methods=['POST'])
def word_to_excel():
    f = request.files['file']
    in_p = os.path.join(UPLOAD_FOLDER, f.filename)
    out_p = in_p.rsplit('.', 1)[0] + '.xlsx'
    f.save(in_p)
    doc = Document(in_p)
    data = []
    for table in doc.tables:
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
    pd.DataFrame(data).to_excel(out_p, index=False, header=False)
    return send_file(out_p, as_attachment=True)

# --- AUTRES ---
@app.route('/img', methods=['POST'])
def img_to_pdf():
    f = request.files['file']
    img = Image.open(f).convert('RGB')
    path = os.path.join(UPLOAD_FOLDER, "image.pdf")
    img.save(path)
    return send_file(path, as_attachment=True)

@app.route('/zip', methods=['POST'])
def to_zip():
    files = request.files.getlist("files")
    path = os.path.join(UPLOAD_FOLDER, "archive.zip")
    with zipfile.ZipFile(path, 'w') as z:
        for f in files:
            fp = os.path.join(UPLOAD_FOLDER, f.filename)
            f.save(fp)
            z.write(fp, f.filename)
    return send_file(path, as_attachment=True)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
